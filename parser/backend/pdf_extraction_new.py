import pdfplumber
# from pypdf2 import PdfReader
import os
import docx
import pytesseract
from PIL import Image
import io
import zipfile
from pdf2image import convert_from_path
import hashlib

RESOLUTION=300
MIN_OCR_CHARS=20
pytesseract.pytesseract.tesseract_cmd = r'C:/Program Files/Tesseract-OCR/tesseract.exe'

# Extracting text from pdfs using pdfplumber
def pdf_extraction(file_path):
  text_content=[]
  seen_lines = set()  # avoid duplicates
  try:
    text_in_resume=""
    with pdfplumber.open(file_path) as pdf:
      pages_as_images = convert_from_path(file_path, dpi=300)
      for i, page in enumerate(pdf.pages):
            # 1. Normal text
            page_text = page.extract_text()
            if page_text:
                for line in page_text.splitlines():
                    clean_line = line.strip()
                    if clean_line and clean_line not in seen_lines:
                        text_in_resume += clean_line + "\n"
                        seen_lines.add(clean_line)

            # 2. OCR text
            ocr_text = pytesseract.image_to_string(pages_as_images[i], lang="eng")
            if ocr_text:
                for line in ocr_text.splitlines():
                    clean_line = line.strip()
                    if clean_line and clean_line not in seen_lines:
                        text_in_resume += clean_line + "\n"
                        seen_lines.add(clean_line)

            # 3. Tables
            tables = page.extract_tables()
            for table in tables:
                for row in table:
                    row_text = " | ".join(cell if cell else "" for cell in row).strip()
                    if row_text and row_text not in seen_lines:
                        text_in_resume += row_text + "\n"
                        seen_lines.add(row_text)
    text_in_resume=text_in_resume.strip()
  except Exception as e:
    print(f"Error processing {file_path}: {e}")
    text_in_resume=""
    with pdfplumber.open(file_path) as resume_pdf:
      for p in resume_pdf.pages:
        text_in_resume=" ".join([text_in_resume,p.extract_text()])
  text_in_resume=text_in_resume.strip()
  return text_in_resume

def extract_text_from_docx(docx_path):
    text_in_resume = ""
    seen_lines = set()
    try:
        # Load document
        doc = docx.Document(docx_path)
        # 1. Normal paragraphs
        for para in doc.paragraphs:
            clean_line = para.text.strip()
            if clean_line and clean_line not in seen_lines:
                text_in_resume += clean_line + "\n"
                seen_lines.add(clean_line)

        # 2. Tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text)
                if row_text and row_text not in seen_lines:
                    text_in_resume += row_text + "\n"
                    seen_lines.add(row_text)

        # 3. OCR from embedded images (DOCX is a zip archive)
        with zipfile.ZipFile(docx_path, "r") as docx_zip:
            for file_name in docx_zip.namelist():
                if file_name.startswith("word/media/"):  # embedded images live here
                    with docx_zip.open(file_name) as image_file:
                        image = Image.open(io.BytesIO(image_file.read()))
                        ocr_text = pytesseract.image_to_string(image, lang="eng")
                        if ocr_text:
                            for line in ocr_text.splitlines():
                                clean_line = line.strip()
                                if clean_line and clean_line not in seen_lines:
                                    text_in_resume += clean_line + "\n"
                                    seen_lines.add(clean_line)
    except Exception as e:
        print(f"Error processing {docx_path}: {e}")
        text_in_resume = ""
    return text_in_resume.strip()

# def extract_text_from_docx(docx_path):
#   doc = docx.Document(docx_path)
#   text = "\n".join([para.text for para in doc.paragraphs])
#   return text

# def table_rows_to_text(tbl, cell_sep=" | ", row_sep="\n"):
#     return row_sep.join(cell_sep.join((c or "").strip() for c in row) for row in tbl).strip()

# def find_tables(page):
#     try:
#         t_objs = page.find_tables()
#     except Exception:
#         return []
#     out = []
#     seen = set()
#     for t in t_objs or []:
#         try:
#             data = t.extract()
#             txt = table_rows_to_text(data)
#             if txt:
#                 h = hashlib.sha1((txt + str(t.bbox)).encode("utf-8")).hexdigest()
#                 if h not in seen:
#                     seen.add(h)
#                     out.append(txt)
#         except Exception:
#             continue
#     return out

# def ocr_page(page):
#     im = page.to_image(resolution=RESOLUTION).original
#     text = pytesseract.image_to_string(im, lang=OCR_LANG).strip()
#     return text if len(text) >= MIN_OCR_CHARS else ""